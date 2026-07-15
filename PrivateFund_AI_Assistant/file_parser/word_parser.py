"""Word 材料解析器。"""

from __future__ import annotations

from pathlib import Path

from docx import Document

from models import SourceChunk


def parse_word(file_path: Path) -> list[SourceChunk]:
    """提取 Word 正文和表格；DOCX 不提供稳定页码，因此使用段落/表格定位。"""
    document = Document(str(file_path))
    chunks: list[SourceChunk] = []
    for index, paragraph in enumerate(document.paragraphs, start=1):
        text = paragraph.text.strip()
        if text:
            chunks.append(SourceChunk(text, file_path.name, f"段落 {index}", "Word"))
    for table_index, table in enumerate(document.tables, start=1):
        rows = [" | ".join(cell.text.strip() for cell in row.cells) for row in table.rows]
        text = "\n".join(row for row in rows if row.strip(" |"))
        if text:
            chunks.append(SourceChunk(text, file_path.name, f"表格 {table_index}", "Word"))
    return chunks

