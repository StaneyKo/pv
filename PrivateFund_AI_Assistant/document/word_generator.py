"""券商研究风格 Word 文件生成器。"""

from __future__ import annotations

import re
from datetime import datetime
from io import BytesIO
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from config.config import OUTPUT_DIR
from document.txt_generator import _safe_filename


def generate_word(company: str, strategy: str, content: str, output_dir: Path = OUTPUT_DIR) -> tuple[bytes, Path]:
    """将结构化文本转换为排版统一的 DOCX。"""
    output_dir.mkdir(parents=True, exist_ok=True)
    document = Document()
    section = document.sections[0]
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    _configure_styles(document)
    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(f"{company}\n{strategy}研究材料")
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(31, 56, 100)
    _set_chinese_font(run, "微软雅黑")

    meta = document.add_paragraph(datetime.now().strftime("生成日期：%Y年%m月%d日"))
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.runs[0].font.size = Pt(9)
    meta.runs[0].font.color.rgb = RGBColor(100, 100, 100)

    for raw_line in content.splitlines():
        line = raw_line.strip()
        if not line:
            continue
        if line.startswith("# 公司名称") or line.startswith("# 策略名称"):
            continue
        if line in {company, strategy}:
            continue
        heading = re.sub(r"^#{1,3}\s*", "", line)
        if line.startswith("##") or re.match(r"^[一二三四五六七八九十]+、", heading):
            document.add_heading(heading, level=1)
        elif line.startswith("#"):
            document.add_heading(heading, level=1)
        elif line.startswith(("- ", "• ")):
            document.add_paragraph(line[2:], style="List Bullet")
        else:
            paragraph = document.add_paragraph(line)
            paragraph.paragraph_format.first_line_indent = Pt(24)
            paragraph.paragraph_format.space_after = Pt(6)
            paragraph.paragraph_format.line_spacing = 1.4

    _add_footer(document)
    buffer = BytesIO()
    document.save(buffer)
    data = buffer.getvalue()
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    path = output_dir / f"{_safe_filename(company)}_{_safe_filename(strategy)}_详细研究_{timestamp}.docx"
    path.write_bytes(data)
    return data, path


def _configure_styles(document: Document) -> None:
    """设置正文与标题的中文字体和颜色。"""
    normal = document.styles["Normal"]
    normal.font.name = "宋体"
    normal.font.size = Pt(10.5)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    heading = document.styles["Heading 1"]
    heading.font.name = "微软雅黑"
    heading.font.size = Pt(13)
    heading.font.bold = True
    heading.font.color.rgb = RGBColor(31, 78, 121)
    heading._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")


def _set_chinese_font(run, font_name: str) -> None:
    """同时设置西文和东亚字体。"""
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)


def _add_footer(document: Document) -> None:
    """添加内部使用提示与自动页码。"""
    footer = document.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("仅供内部研究使用  |  第 ")
    run.font.size = Pt(8)
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    footer._p.append(field)
    footer.add_run(" 页")

